"""
insert_report_images.py — 把截图按序插入实习报告对应章节，并加图注。
基于已做文字修正的 生产实习报告_完成版.docx（会先备份为 _插图前备份）。
"""
import shutil
from pathlib import Path
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

DESK = Path(r"C:\Users\32698\Desktop")
# 从无图的文字修正版读取，输出到项目目录（避开 WPS Cloud 同步锁）
BASE = DESK / "生产实习报告_完成版_插图前备份.docx"
OUT = Path(__file__).parent.parent / "生产实习报告_完成版_带图.docx"
IMG = DESK / "报告截图"

doc = docx.Document(str(BASE))
SRC = OUT  # 保存目标

# 锚点关键词 -> 该锚点后要插入的 [(图文件, 图注), ...]
PLAN = [
    ("本项目采用Pipeline RAG双分支工作流架构，共8个节点", [
        ("1.png", "图4-1  Dify Workflow 双分支工作流全景（知识库命中 / 联网兜底）"),
    ]),
    ("知识库覆盖13大类设备，97个文档（92篇Markdown", [
        ("2.png", "图4-2  知识库文档列表（父子索引库，97个文档）"),
        ("3.png", "图4-3  召回测试：全文父块模式返回完整正文（含核心知识/可能原因/检查步骤）"),
    ]),
    ("项目部署环境基于Docker Desktop", [
        ("6.png", "图5-1  Docker 容器运行状态（14个服务全部 Running）"),
    ]),
    ("3个典型场景人工演示验证", [
        ("4-1.png", "图6-1  场景一（知识库命中）：输入设备故障，返回结构化排查步骤"),
        ("4-2.png", "图6-2  场景一续：答案附带知识库来源引用"),
        ("5-1.png", "图6-3  场景二（联网兜底）：知识库未命中，自动触发 Tavily 搜索"),
        ("5-2.png", "图6-4  场景二续：基于联网结果生成诊断建议并标注 URL 来源"),
    ]),
]


def find_para(keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None


BODY_PARENT = doc.paragraphs[0]._parent  # 有 .part，图片嵌入需要


def insert_after(ref_elem, img_path, caption):
    """在 ref_elem 之后插入 图片段 + 图注段，返回图注段元素（供链式插入）。"""
    # 图片段
    img_p = OxmlElement("w:p")
    ref_elem.addnext(img_p)
    para_img = docx.text.paragraph.Paragraph(img_p, BODY_PARENT)
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para_img.add_run()
    run.add_picture(str(img_path), width=Inches(6.0))

    # 图注段
    cap_p = OxmlElement("w:p")
    img_p.addnext(cap_p)
    para_cap = docx.text.paragraph.Paragraph(cap_p, BODY_PARENT)
    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para_cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.name = "宋体"
    return cap_p


total = 0
for keyword, imgs in PLAN:
    anchor = find_para(keyword)
    if anchor is None:
        print(f"[跳过] 未找到锚点: {keyword[:20]}")
        continue
    cursor = anchor._p
    for fname, caption in imgs:
        fpath = IMG / fname
        if not fpath.exists():
            print(f"  [缺图] {fname}")
            continue
        cursor = insert_after(cursor, fpath, caption)
        total += 1
        print(f"  [插入] {fname} → {keyword[:14]}...")

doc.save(str(SRC))
print(f"\n完成：插入 {total} 张图")
