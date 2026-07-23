"""
insert_team_table.py — 在报告体会部分插入「项目组成员分工与贡献」表格。
数据来源：33组 PPT 第13/14页团队分工。
作用于 生产实习报告_完成版_带图.docx（项目目录）。
"""
from pathlib import Path
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

SRC = Path(__file__).parent.parent / "生产实习报告_完成版_带图.docx"
doc = docx.Document(str(SRC))
BODY_PARENT = doc.paragraphs[0]._parent

# 表格数据
HEADER = ["成员", "主要职责", "具体贡献"]
ROWS = [
    ["肖  扬", "基础框架搭建 · 平台运维",
     "负责 Dify 平台基础框架搭建与技术文档书写，保障整套系统稳定上线运行。"],
    ["杨金浩", "基础框架 · 文档书写",
     "与肖扬搭档完成基础框架与文档书写，持续维护平台基本运行与环境调优。"],
    ["王  帅", "功能完善 · 编辑优化",
     "在基础框架之上完善功能细节，添加并打磨问答、引用、多轮追问等编辑型能力。"],
    ["徐  熠", "调研 · 排查 · 汇报",
     "负责前期技术调研、阶段性问题排查定位，并统筹答辩材料的整理与制作。"],
]

# 找锚点段落
anchor = None
for p in doc.paragraphs:
    if "团队协作方面，四名成员分工明确" in p.text:
        anchor = p
        break
if anchor is None:
    raise SystemExit("未找到分工锚点段落")

# 建表（先加到文末，再移动到锚点后）
table = doc.add_table(rows=1 + len(ROWS), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
for j, h in enumerate(HEADER):
    cell = table.rows[0].cells[j]
    cell.text = ""
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(10.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 数据行
for i, row in enumerate(ROWS, start=1):
    for j, val in enumerate(row):
        cell = table.rows[i].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(10.5)
        if j < 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表标题段（插在锚点后，表在标题后）
title_p = OxmlElement("w:p")
anchor._p.addnext(title_p)
para_title = docx.text.paragraph.Paragraph(title_p, BODY_PARENT)
para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = para_title.add_run("表1  项目组（第33组）成员分工与贡献")
rt.bold = True
rt.font.size = Pt(10.5)

# 把表移动到标题段之后
title_p.addnext(table._tbl)

doc.save(str(SRC))
print("完成：已插入分工贡献表格（4名成员 + 表头）")
