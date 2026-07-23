"""
fix_report.py — 修正生产实习报告的技术参数偏差，使其与最终交付版本一致。
保留原 docx 格式，仅逐段替换文本。
"""
import docx
import shutil
from pathlib import Path

SRC = Path(r"C:\Users\32698\Desktop\生产实习报告_完成版.docx")
BAK = SRC.with_name("生产实习报告_完成版_修改前备份.docx")

# 先备份
if not BAK.exists():
    shutil.copy2(SRC, BAK)
    print(f"已备份: {BAK.name}")

doc = docx.Document(str(SRC))

# 整段替换映射：段落索引(以 report.txt 的编号) -> 新全文
# 用「包含关键词匹配」而非硬索引，避免段落漂移
REPLACEMENTS = [
    # (匹配用的旧片段, 新的整段文本)
    (
        "外部服务层：DeepSeek-V4-Pro",
        "外部服务层：DeepSeek-V4-Pro（经SiliconFlow调用）大语言模型、bge-large-zh-v1.5中文Embedding向量化模型、bge-reranker-v2-m3重排模型、Tavily联网搜索API；",
    ),
    (
        "数据层：预置知识库包含13大类设备",
        "数据层：预置知识库包含13大类设备、97个文档（92篇Markdown + 5篇PDF手册）。",
    ),
    (
        "系统关键流程为：自然语言提问",
        "系统关键流程为：自然语言提问 → 工作流编排 → 知识库检索（父子索引库，top_k=8，启用score_threshold=0.4阈值过滤，开启rerank） → 未命中（最高分<0.4）则联网搜索 → LLM综合生成 → 返回结构化答案及引用。整个系统部署在一台本地机器上，外部仅依赖两个API：SiliconFlow（LLM）与Tavily（联网搜索）。",
    ),
    (
        "本项目采用Pipeline RAG双分支工作流架构，共6个节点",
        "本项目采用Pipeline RAG双分支工作流架构，共8个节点：Start节点接收用户提问；知识库检索节点（父子索引库，top_k=8，启用score_threshold=0.4阈值过滤，开启rerank）；条件判断节点（检索结果非空且最高分≥0.4判定为命中，否则未命中）；Tavily搜索节点（max_results=5，仅未命中时触发）；两条路径各自独立的LLM综合节点（知识库路径llm_kb，temperature=0.3；联网路径llm_generate，temperature=0.6）；两条路径各自独立的回复输出节点（answer_kb与answer）。工作流设计要点为两条分支各自独立，便于知识库命中与联网场景的LLM参数独立调优；固定流程优先于「LLM自主决策」，确保演示稳定性。",
    ),
    (
        "知识库覆盖13大类设备，约102个文档",
        "知识库覆盖13大类设备，97个文档（92篇Markdown + 5篇PDF手册）。切片策略采用Dify父子索引（Parent-Child）模式并将父块设为「全文」——子块用于精准语义匹配，命中后返回整篇文档作为父块上下文。此设计解决了通用分段模式下「泛泛查询只命中文档标题元数据、拿不到正文」的问题：即便用户提问较宽泛（如「除尘常见问题」），也能召回包含核心知识、可能原因、检查步骤的完整正文。向量化采用bge-large-zh-v1.5中文模型（1024维），存储在Weaviate向量数据库中。混合检索策略为私有知识库优先（BGE-zh Embedding → Weaviate向量库 + Rerank提升精度），当知识库检索结果最高分低于0.4阈值时自动触发Tavily联网搜索作为兜底。",
    ),
    (
        "创建知识库「制造业设备维修手册库」",
        "创建知识库「制造业设备维修-父子索引」，分段模式选择Parent-Child（父子索引）、父块模式设为「全文」，预上传92篇Markdown设备分类文档及5篇补充PDF手册（共97个），并启动高质量索引。为避免不同设备类别目录下同名Markdown文件（如各类别均有04_常见故障与可能原因.md）被Dify按文件名去重覆盖，上传前使用make_unique_named_copy.py脚本为文件加「类别目录名」前缀生成唯一命名副本。",
    ),
    (
        "通过Dify控制台「工作室」→「导入DSL文件」方式导入预配置的chatflow-dsl.yml",
        "通过Dify控制台「工作室」→「导入DSL文件」方式导入预配置的chatflow-dsl.yml。导入后在Workflow编辑器中关联实际父子索引知识库（替换DATASET_ID），在知识库检索节点打开「Score阈值」开关并设为0.4（此开关不打开则阈值不生效，会导致弱相关文档误命中、挡住联网兜底分支），配置tavily_search HTTP Request节点填入真实Tavily密钥（替换REPLACE_TAVILY_KEY），粘贴完整System Prompt到LLM节点。调试时重点关注条件分支判断逻辑和双分支LLM温度的独立调优，最终发布为Chatbot应用。",
    ),
    (
        "repair_dataset.py用于修复索引错误文档",
        "项目配套开发了Python自动化脚本：make_unique_named_copy.py用于生成唯一命名副本以规避Dify同名去重；rebuild_parent_kb_fulldoc.py用于以全文父块模式重建父子索引库；upload_to_both_kb.py用于向通用库与父子库并发上传；sync_knowledge_base.py用于新手册放入data/new-manuals目录后的自动上传与索引状态监控；eval_chatflow.py用于运行20题自动评测，评估关键词覆盖率、引用来源标注率、流式首字响应时间等指标。",
    ),
]

changed = 0
for p in doc.paragraphs:
    txt = p.text
    for old_frag, new_full in REPLACEMENTS:
        if old_frag in txt:
            # 清空该段所有 run，把新文本写入第一个 run（保留段落样式）
            if p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new_full)
            changed += 1
            print(f"[改] {old_frag[:20]}...")
            break

doc.save(str(SRC))
print(f"\n完成：修改了 {changed} 段")
