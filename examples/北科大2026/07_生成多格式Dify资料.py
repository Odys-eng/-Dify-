from __future__ import annotations

import csv
import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "练习3资料"
MARKDOWN_FILE = OUTPUT_DIR / "Dify学习资料汇编.md"
BASE_NAME = "Dify学习资料汇编"

ROADMAP = [
    ["第 1 周", "环境、模型与应用类型", "可访问的本地 Dify；一个 Chatbot", "Web/API 可访问，模型调用正常"],
    ["第 2 周", "知识库与 RAG", "一个包含多格式资料的知识库", "文档索引 completed，问题可召回"],
    ["第 3 周", "Workflow 与 Chatflow", "知识检索到 LLM 的问答流程", "有正确回答、引用和无结果兜底"],
    ["第 4 周", "Agent、API 与评估", "API 脚本、检索测试集和调优记录", "可重复运行并解释参数效果"],
]


def markdown_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    in_frontmatter = False
    for raw in text.splitlines():
        line = raw.strip()
        if line == "---":
            in_frontmatter = not in_frontmatter
            continue
        if in_frontmatter or not line or line.startswith("|"):
            continue
        if line.startswith("# "):
            blocks.append(("title", line[2:]))
        elif line.startswith("## "):
            blocks.append(("heading", line[3:]))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:]))
        elif re.match(r"^\d+\. ", line):
            blocks.append(("number", re.sub(r"^\d+\. ", "", line)))
        else:
            blocks.append(("paragraph", line))
    return blocks


def plain_text(text: str) -> str:
    text = re.sub(r"^---.*?---\s*", "", text, flags=re.DOTALL)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*|`([^`]+)`", lambda m: m.group(1) or m.group(2), text)
    return text


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def build_docx(text: str, output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for style_name, size, color in (("Title", 26, "17324D"), ("Heading 1", 17, "1F5C78")):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "RuyiDify 学习资料 | 2026-07-18"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Dify 官方资料学习汇编")
    run.font.size = Pt(8)

    for kind, value in markdown_blocks(text):
        if kind == "title":
            p = document.add_paragraph(value, style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph("基于 Dify 官方文档整理 | 更新日期 2026-07-18").alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "heading":
            document.add_heading(value, level=1)
        elif kind == "bullet":
            document.add_paragraph(value, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(value, style="List Number")
        elif not value.startswith("资料核对日期"):
            document.add_paragraph(value.replace("`", ""))

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录：四周学习计划", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["周次", "主题", "实践产出", "验证标准"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
        set_cell_shading(table.rows[0].cells[index], "D9EAF0")
    for row in ROADMAP:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.save(output)


def build_pdf(text: str, output: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("CJKTitle", parent=base["Title"], fontName="STSong-Light", fontSize=24, leading=32, alignment=TA_CENTER, textColor=colors.HexColor("#17324D"), spaceAfter=18),
        "heading": ParagraphStyle("CJKHeading", parent=base["Heading1"], fontName="STSong-Light", fontSize=15, leading=21, textColor=colors.HexColor("#1F5C78"), spaceBefore=12, spaceAfter=7),
        "body": ParagraphStyle("CJKBody", parent=base["BodyText"], fontName="STSong-Light", fontSize=9.5, leading=15, textColor=colors.HexColor("#263238"), spaceAfter=6),
    }
    story = []
    for kind, value in markdown_blocks(text):
        safe = html.escape(value.replace("`", ""))
        if kind == "title":
            story.append(Paragraph(safe, styles["title"]))
            story.append(Paragraph("基于 Dify 官方文档整理 | 更新日期 2026-07-18", styles["body"]))
            story.append(Spacer(1, 10))
        elif kind == "heading":
            story.append(Paragraph(safe, styles["heading"]))
        elif kind == "bullet":
            story.append(Paragraph("• " + safe, styles["body"]))
        elif kind == "number":
            story.append(Paragraph(safe, styles["body"]))
        else:
            story.append(Paragraph(safe, styles["body"]))
    story.append(PageBreak())
    story.append(Paragraph("附录：四周学习计划", styles["heading"]))
    data = [["周次", "主题", "实践产出", "验证标准"], *ROADMAP]
    table = Table(data, colWidths=[1.6 * cm, 3.2 * cm, 5.2 * cm, 5.2 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF0")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94A3B8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEADING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    document = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=1.8 * cm, leftMargin=1.8 * cm, topMargin=1.8 * cm, bottomMargin=1.8 * cm, title=BASE_NAME, author="RuyiDify")
    document.build(story)


def build_html(text: str, output: Path) -> None:
    parts = ["<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>", "<title>Dify学习资料汇编</title>", "<style>body{font-family:'Microsoft YaHei',sans-serif;max-width:900px;margin:40px auto;line-height:1.75;color:#263238}h1{color:#17324d}h2{color:#1f5c78;border-bottom:1px solid #d9eaf0}li{margin:4px 0}</style></head><body>"]
    list_open = False
    for kind, value in markdown_blocks(text):
        safe = html.escape(value.replace("`", ""))
        if kind in {"bullet", "number"}:
            if not list_open:
                parts.append("<ul>")
                list_open = True
            parts.append(f"<li>{safe}</li>")
            continue
        if list_open:
            parts.append("</ul>")
            list_open = False
        tag = {"title": "h1", "heading": "h2"}.get(kind, "p")
        parts.append(f"<{tag}>{safe}</{tag}>")
    if list_open:
        parts.append("</ul>")
    parts.append("</body></html>")
    output.write_text("\n".join(parts), encoding="utf-8")


def build_tabular_files() -> None:
    csv_file = OUTPUT_DIR / f"{BASE_NAME}_四周路线.csv"
    with csv_file.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["周次", "主题", "实践产出", "验证标准"])
        writer.writerows(ROADMAP)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Dify四周学习路线"
    sheet.append(["周次", "主题", "实践产出", "验证标准"])
    for row in ROADMAP:
        sheet.append(row)
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="17324D")
        cell.fill = PatternFill("solid", fgColor="D9EAF0")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    widths = [12, 28, 42, 42]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    workbook.save(OUTPUT_DIR / f"{BASE_NAME}_四周路线.xlsx")


def main() -> None:
    text = MARKDOWN_FILE.read_text(encoding="utf-8")
    (OUTPUT_DIR / f"{BASE_NAME}.txt").write_text(plain_text(text), encoding="utf-8")
    build_html(text, OUTPUT_DIR / f"{BASE_NAME}.html")
    build_docx(text, OUTPUT_DIR / f"{BASE_NAME}.docx")
    build_pdf(text, OUTPUT_DIR / f"{BASE_NAME}.pdf")
    build_tabular_files()
    print("已生成 6 种格式：MD、DOCX、PDF、TXT、HTML、CSV、XLSX（MD 为源文件，共 7 个文件）")
    for path in sorted(OUTPUT_DIR.glob(f"{BASE_NAME}*")):
        print(f"{path.name}: {path.stat().st_size} bytes")


if __name__ == "__main__":
    main()
